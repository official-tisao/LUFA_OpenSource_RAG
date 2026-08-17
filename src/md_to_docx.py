#!/usr/bin/env python3
"""
md_to_docx.py: convert a thesis chapter written in Markdown into a Word document.

Deliberately narrow. It handles only the constructs the thesis chapters actually use, namely
ATX headings, paragraphs, bullet lists, pipe tables, blockquotes, fenced code blocks, and
inline bold/italic/code.
A general Markdown engine would pull in a dependency and still need the same styling work.

Styling follows the project default: Times New Roman 12 pt, 1.5 line spacing, justified body
text. Headings are Times New Roman too, so the document does not mix families.

Usage:
  python src/md_to_docx.py thesis/Chapter6_Discussion.md thesis/Chapter6.docx
  python src/md_to_docx.py in.md out.docx --title "Chapter Six"
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

FONT = "Times New Roman"
BODY_PT = 12
HEADING_PT = {1: 16, 2: 14, 3: 12.5, 4: 12}

# Inline spans, longest delimiter first so ** is not eaten by *.
_INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*.+?\*)", re.DOTALL)


def _set_font(run, size=BODY_PT, bold=False, italic=False, mono=False, color=None):
    run.font.name = "Consolas" if mono else FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # python-docx sets only the latin typeface; east-asian must be set on the XML directly or
    # Word substitutes a default font for any character it classifies as non-latin.
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rf)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(attr), "Consolas" if mono else FONT)


def _add_inline(par, text, size=BODY_PT, bold=False, italic=False, color=None):
    """Render inline **bold**, *italic* and `code` into an existing paragraph."""
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            _set_font(par.add_run(part[2:-2]), size, True, italic, color=color)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            _set_font(par.add_run(part[1:-1]), size - 1, bold, italic, mono=True, color=color)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            _set_font(par.add_run(part[1:-1]), size, bold, True, color=color)
        else:
            _set_font(par.add_run(part), size, bold, italic, color=color)


def _is_table_sep(line):
    s = line.strip()
    return bool(s) and set(s) <= set("|-: ") and "-" in s and "|" in s


def _cells(line):
    s = line.strip().strip("|")
    return [c.strip() for c in s.split("|")]


def convert(md_path, docx_path, title=None):
    lines = Path(md_path).read_text(encoding="utf-8").splitlines()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(BODY_PT)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(title), HEADING_PT[1], bold=True)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Fenced code block. Rendered verbatim, one paragraph per line, left aligned:
        # justification would stretch the spacing inside shell commands.
        if stripped.startswith("```"):
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.left_indent = Pt(18)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                _set_font(p.add_run(lines[i].rstrip()), size=10, mono=True)
                i += 1
            i += 1  # closing fence
            doc.add_paragraph()
            continue

        # Pipe table: a header row followed by a separator row.
        if stripped.startswith("|") and i + 1 < n and _is_table_sep(lines[i + 1]):
            header = _cells(stripped)
            # Wide tables overflow the portrait text block at the usual 10.5 pt.
            cell_pt = 9 if len(header) >= 7 else 10.5
            i += 2
            body = []
            while i < n and lines[i].strip().startswith("|"):
                body.append(_cells(lines[i].strip()))
                i += 1
            t = doc.add_table(rows=1, cols=len(header))
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.autofit = True
            for c, htxt in zip(t.rows[0].cells, header):
                c.paragraphs[0].text = ""
                _add_inline(c.paragraphs[0], htxt, size=cell_pt, bold=True)
            for row in body:
                cells = t.add_row().cells
                # A malformed row must not raise: pad or truncate to the header width.
                for c, txt in zip(cells, row + [""] * (len(header) - len(row))):
                    c.paragraphs[0].text = ""
                    _add_inline(c.paragraphs[0], txt, size=cell_pt)
            doc.add_paragraph()
            continue

        # Horizontal rule.
        if re.fullmatch(r"-{3,}|_{3,}|\*{3,}", stripped):
            i += 1
            continue

        # Heading.
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            lvl = min(len(m.group(1)), 4)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14 if lvl <= 2 else 10)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            _add_inline(p, m.group(2), size=HEADING_PT[lvl], bold=True)
            i += 1
            continue

        # Blockquote: consume the whole run, render indented and italic grey.
        if stripped.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            for chunk in re.split(r"\n\s*\n", "\n".join(buf).strip()):
                if not chunk.strip():
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(24)
                p.paragraph_format.space_after = Pt(6)
                _add_inline(p, " ".join(chunk.split()), size=11,
                            italic=True, color=RGBColor(0x44, 0x44, 0x44))
            continue

        # Bullet or numbered list item, joining any wrapped continuation lines.
        m = re.match(r"^([-*+]|\d+\.)\s+(.*)$", stripped)
        if m:
            text = m.group(2)
            i += 1
            while i < n and lines[i].strip() and not re.match(
                    r"^(#{1,6}\s|>|[-*+]\s|\d+\.\s|```|\|)", lines[i].strip()) \
                    and lines[i].startswith((" ", "\t")):
                text += " " + lines[i].strip()
                i += 1
            p = doc.add_paragraph(style="List Number" if m.group(1)[0].isdigit()
                                  else "List Bullet")
            p.paragraph_format.space_after = Pt(4)
            _add_inline(p, text)
            continue

        # Paragraph: join wrapped lines until a blank line or a new block construct.
        buf = [stripped]
        i += 1
        while i < n and lines[i].strip() and not re.match(
                r"^(#{1,6}\s|>|[-*+]\s|\d+\.\s|```|\||-{3,}$)", lines[i].strip()):
            buf.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_inline(p, " ".join(buf))

    Path(docx_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print(f"[docx] wrote {docx_path}")
    return 0


def main():
    ap = argparse.ArgumentParser(description="Convert a Markdown chapter to .docx.")
    ap.add_argument("md")
    ap.add_argument("docx")
    ap.add_argument("--title")
    a = ap.parse_args()
    if not Path(a.md).exists():
        print(f"[docx] not found: {a.md}")
        return 1
    return convert(a.md, a.docx, a.title)


if __name__ == "__main__":
    sys.exit(main())
